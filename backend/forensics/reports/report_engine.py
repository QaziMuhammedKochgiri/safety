"""
Report Engine for SafeChild
Court-ready report generation with PDF, DOCX support and digital signatures.
"""

import hashlib
import os
import io
import base64
from datetime import datetime
from enum import Enum
from typing import Optional, List, Dict, Any, Union
from dataclasses import dataclass, field
from pathlib import Path
import json
import logging

logger = logging.getLogger(__name__)


class ReportType(Enum):
    """Types of forensic reports."""
    FULL_ANALYSIS = "full_analysis"
    EXECUTIVE_SUMMARY = "executive_summary"
    EVIDENCE_INDEX = "evidence_index"
    CHAIN_OF_CUSTODY = "chain_of_custody"
    EXPERT_WITNESS = "expert_witness"
    RISK_ASSESSMENT = "risk_assessment"
    ALIENATION_REPORT = "alienation_report"
    TIMELINE = "timeline"


class ReportFormat(Enum):
    """Output formats for reports."""
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    JSON = "json"
    E001 = "e001"
    CELLEBRITE_XML = "cellebrite_xml"


class ReportLanguage(Enum):
    """Supported report languages."""
    ENGLISH = "en"
    GERMAN = "de"
    TURKISH = "tr"


@dataclass
class SignatureInfo:
    """Digital signature information."""
    signer_name: str
    signer_title: str
    signer_organization: str
    signature_date: datetime
    certificate_serial: Optional[str] = None
    signature_hash: Optional[str] = None
    verification_url: Optional[str] = None


@dataclass
class EvidenceItem:
    """Individual evidence item for reports."""
    item_id: str
    item_type: str  # message, image, call_log, etc.
    source: str     # WhatsApp, Instagram, etc.
    timestamp: datetime
    description: str
    hash_value: str
    file_path: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    risk_indicators: List[str] = field(default_factory=list)
    relevance_score: float = 0.0


@dataclass
class ChainOfCustody:
    """Chain of custody tracking."""
    case_id: str
    evidence_id: str
    acquisition_date: datetime
    acquisition_method: str
    acquired_by: str
    device_info: Dict[str, str]
    hash_original: str
    hash_algorithm: str = "SHA-256"
    custody_log: List[Dict[str, Any]] = field(default_factory=list)

    def add_custody_entry(self, action: str, person: str, notes: str = ""):
        """Add entry to custody log."""
        entry = {
            "timestamp": datetime.utcnow().isoformat(),
            "action": action,
            "person": person,
            "notes": notes
        }
        self.custody_log.append(entry)
        return entry


@dataclass
class GeneratedReport:
    """Result of report generation."""
    report_id: str
    report_type: ReportType
    format: ReportFormat
    language: ReportLanguage
    title: str
    generated_at: datetime
    content: bytes
    content_hash: str
    page_count: int
    signature: Optional[SignatureInfo] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class ReportTemplate:
    """Base class for report templates."""

    # Translations for common terms
    TRANSLATIONS = {
        "en": {
            "report_title": "Forensic Analysis Report",
            "case_number": "Case Number",
            "date": "Date",
            "prepared_by": "Prepared By",
            "confidential": "CONFIDENTIAL - ATTORNEY-CLIENT PRIVILEGED",
            "executive_summary": "Executive Summary",
            "evidence_overview": "Evidence Overview",
            "key_findings": "Key Findings",
            "risk_assessment": "Risk Assessment",
            "chain_of_custody": "Chain of Custody",
            "methodology": "Methodology",
            "conclusions": "Conclusions",
            "recommendations": "Recommendations",
            "appendix": "Appendix",
            "page": "Page",
            "of": "of",
            "table_of_contents": "Table of Contents",
            "digital_signature": "Digital Signature",
            "hash_verification": "Hash Verification",
            "evidence_item": "Evidence Item",
            "timestamp": "Timestamp",
            "source": "Source",
            "description": "Description",
            "risk_level": "Risk Level",
            "high": "HIGH",
            "medium": "MEDIUM",
            "low": "LOW",
        },
        "de": {
            "report_title": "Forensischer Analysebericht",
            "case_number": "Aktenzeichen",
            "date": "Datum",
            "prepared_by": "Erstellt von",
            "confidential": "VERTRAULICH - ANWALTSGEHEIMNIS",
            "executive_summary": "Zusammenfassung",
            "evidence_overview": "Beweisübersicht",
            "key_findings": "Wichtigste Erkenntnisse",
            "risk_assessment": "Risikobewertung",
            "chain_of_custody": "Beweismittelkette",
            "methodology": "Methodik",
            "conclusions": "Schlussfolgerungen",
            "recommendations": "Empfehlungen",
            "appendix": "Anhang",
            "page": "Seite",
            "of": "von",
            "table_of_contents": "Inhaltsverzeichnis",
            "digital_signature": "Digitale Signatur",
            "hash_verification": "Hash-Verifizierung",
            "evidence_item": "Beweismittel",
            "timestamp": "Zeitstempel",
            "source": "Quelle",
            "description": "Beschreibung",
            "risk_level": "Risikostufe",
            "high": "HOCH",
            "medium": "MITTEL",
            "low": "NIEDRIG",
        },
        "tr": {
            "report_title": "Adli Analiz Raporu",
            "case_number": "Dosya Numarasi",
            "date": "Tarih",
            "prepared_by": "Hazirlayan",
            "confidential": "GIZLI - AVUKAT-MUVEKKIL AYRICALIKLI",
            "executive_summary": "Yonetici Ozeti",
            "evidence_overview": "Delil Genel Bakis",
            "key_findings": "Temel Bulgular",
            "risk_assessment": "Risk Degerlendirmesi",
            "chain_of_custody": "Delil Zinciri",
            "methodology": "Metodoloji",
            "conclusions": "Sonuclar",
            "recommendations": "Oneriler",
            "appendix": "Ek",
            "page": "Sayfa",
            "of": "/",
            "table_of_contents": "Icindekiler",
            "digital_signature": "Dijital Imza",
            "hash_verification": "Hash Dogrulama",
            "evidence_item": "Delil",
            "timestamp": "Zaman Damgasi",
            "source": "Kaynak",
            "description": "Aciklama",
            "risk_level": "Risk Seviyesi",
            "high": "YUKSEK",
            "medium": "ORTA",
            "low": "DUSUK",
        }
    }

    def __init__(self, language: ReportLanguage = ReportLanguage.ENGLISH):
        self.language = language
        self.strings = self.TRANSLATIONS.get(language.value, self.TRANSLATIONS["en"])

    def t(self, key: str) -> str:
        """Get translated string."""
        return self.strings.get(key, key)

    def get_css_styles(self) -> str:
        """Return CSS styles for HTML/PDF reports."""
        return """
        @page {
            size: A4;
            margin: 2cm;
            @top-center {
                content: "CONFIDENTIAL";
                font-size: 10pt;
                color: #cc0000;
            }
            @bottom-center {
                content: "Page " counter(page) " of " counter(pages);
                font-size: 10pt;
            }
        }

        body {
            font-family: 'DejaVu Sans', 'Noto Sans', Arial, sans-serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
        }

        .header {
            text-align: center;
            border-bottom: 2px solid #1a365d;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }

        .header h1 {
            color: #1a365d;
            font-size: 24pt;
            margin-bottom: 10px;
        }

        .confidential-banner {
            background-color: #cc0000;
            color: white;
            text-align: center;
            padding: 10px;
            font-weight: bold;
            margin-bottom: 20px;
        }

        .case-info {
            background-color: #f7fafc;
            border: 1px solid #e2e8f0;
            padding: 15px;
            margin-bottom: 20px;
        }

        .case-info table {
            width: 100%;
        }

        .case-info td {
            padding: 5px 10px;
        }

        .case-info td:first-child {
            font-weight: bold;
            width: 200px;
        }

        h2 {
            color: #1a365d;
            border-bottom: 1px solid #e2e8f0;
            padding-bottom: 10px;
            margin-top: 30px;
        }

        h3 {
            color: #2d3748;
            margin-top: 20px;
        }

        .evidence-table {
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }

        .evidence-table th {
            background-color: #1a365d;
            color: white;
            padding: 10px;
            text-align: left;
        }

        .evidence-table td {
            border: 1px solid #e2e8f0;
            padding: 10px;
        }

        .evidence-table tr:nth-child(even) {
            background-color: #f7fafc;
        }

        .risk-high {
            color: #c53030;
            font-weight: bold;
        }

        .risk-medium {
            color: #dd6b20;
            font-weight: bold;
        }

        .risk-low {
            color: #38a169;
            font-weight: bold;
        }

        .chain-of-custody {
            border: 2px solid #1a365d;
            padding: 20px;
            margin: 20px 0;
        }

        .chain-of-custody h3 {
            margin-top: 0;
        }

        .signature-block {
            margin-top: 50px;
            border-top: 1px solid #e2e8f0;
            padding-top: 20px;
        }

        .signature-line {
            border-bottom: 1px solid #333;
            width: 300px;
            margin: 30px 0 5px 0;
        }

        .hash-verification {
            font-family: monospace;
            background-color: #f7fafc;
            padding: 15px;
            border: 1px solid #e2e8f0;
            word-break: break-all;
        }

        .toc {
            margin: 30px 0;
        }

        .toc a {
            color: #1a365d;
            text-decoration: none;
        }

        .toc ul {
            list-style-type: none;
            padding-left: 0;
        }

        .toc li {
            padding: 5px 0;
            border-bottom: 1px dotted #e2e8f0;
        }

        .timeline {
            position: relative;
            padding-left: 30px;
        }

        .timeline::before {
            content: '';
            position: absolute;
            left: 10px;
            top: 0;
            bottom: 0;
            width: 2px;
            background-color: #1a365d;
        }

        .timeline-item {
            position: relative;
            margin-bottom: 20px;
            padding-left: 20px;
        }

        .timeline-item::before {
            content: '';
            position: absolute;
            left: -24px;
            top: 5px;
            width: 10px;
            height: 10px;
            border-radius: 50%;
            background-color: #1a365d;
        }

        .timeline-date {
            font-weight: bold;
            color: #1a365d;
        }

        .appendix {
            page-break-before: always;
        }
        """


class PDFGenerator:
    """Generate PDF reports using WeasyPrint."""

    def __init__(self):
        self._weasyprint = None

    def _load_weasyprint(self):
        """Lazy load WeasyPrint."""
        if self._weasyprint is None:
            try:
                from weasyprint import HTML, CSS
                self._weasyprint = {"HTML": HTML, "CSS": CSS}
                logger.info("WeasyPrint loaded successfully")
            except ImportError:
                logger.warning("WeasyPrint not available, PDF generation disabled")
                self._weasyprint = {}
        return self._weasyprint

    def generate(self, html_content: str, css_styles: str = "") -> Optional[bytes]:
        """Generate PDF from HTML content."""
        wp = self._load_weasyprint()
        if not wp:
            logger.error("Cannot generate PDF: WeasyPrint not installed")
            return None

        try:
            html = wp["HTML"](string=html_content)
            css = wp["CSS"](string=css_styles) if css_styles else None

            pdf_buffer = io.BytesIO()
            if css:
                html.write_pdf(pdf_buffer, stylesheets=[css])
            else:
                html.write_pdf(pdf_buffer)

            pdf_bytes = pdf_buffer.getvalue()
            logger.info(f"Generated PDF: {len(pdf_bytes)} bytes")
            return pdf_bytes

        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return None

    def get_page_count(self, html_content: str, css_styles: str = "") -> int:
        """Get page count for generated PDF."""
        wp = self._load_weasyprint()
        if not wp:
            return 0

        try:
            html = wp["HTML"](string=html_content)
            css = wp["CSS"](string=css_styles) if css_styles else None

            if css:
                doc = html.render(stylesheets=[css])
            else:
                doc = html.render()

            return len(doc.pages)
        except Exception as e:
            logger.error(f"Page count failed: {e}")
            return 0


class DOCXGenerator:
    """Generate DOCX reports using python-docx."""

    def __init__(self):
        self._docx = None

    def _load_docx(self):
        """Lazy load python-docx."""
        if self._docx is None:
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.style import WD_STYLE_TYPE
                self._docx = {
                    "Document": Document,
                    "Inches": Inches,
                    "Pt": Pt,
                    "RGBColor": RGBColor,
                    "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
                    "WD_STYLE_TYPE": WD_STYLE_TYPE
                }
                logger.info("python-docx loaded successfully")
            except ImportError:
                logger.warning("python-docx not available, DOCX generation disabled")
                self._docx = {}
        return self._docx

    def generate(self, report_data: Dict[str, Any], template: ReportTemplate) -> Optional[bytes]:
        """Generate DOCX from report data."""
        docx = self._load_docx()
        if not docx:
            logger.error("Cannot generate DOCX: python-docx not installed")
            return None

        try:
            doc = docx["Document"]()

            # Add title
            title = doc.add_heading(report_data.get("title", template.t("report_title")), 0)
            title.alignment = docx["WD_ALIGN_PARAGRAPH"].CENTER

            # Add confidential notice
            p = doc.add_paragraph()
            p.alignment = docx["WD_ALIGN_PARAGRAPH"].CENTER
            run = p.add_run(template.t("confidential"))
            run.bold = True
            run.font.color.rgb = docx["RGBColor"](204, 0, 0)

            # Add case info
            doc.add_heading(template.t("case_number"), level=1)
            case_info = report_data.get("case_info", {})
            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"

            for key, value in case_info.items():
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = str(value)

            # Add sections
            for section in report_data.get("sections", []):
                doc.add_heading(section.get("title", ""), level=1)
                doc.add_paragraph(section.get("content", ""))

            # Add evidence table
            if "evidence_items" in report_data:
                doc.add_heading(template.t("evidence_overview"), level=1)
                evidence_table = doc.add_table(rows=1, cols=5)
                evidence_table.style = "Table Grid"

                headers = evidence_table.rows[0].cells
                headers[0].text = "ID"
                headers[1].text = template.t("timestamp")
                headers[2].text = template.t("source")
                headers[3].text = template.t("description")
                headers[4].text = template.t("risk_level")

                for item in report_data["evidence_items"]:
                    row = evidence_table.add_row().cells
                    row[0].text = item.get("item_id", "")
                    row[1].text = str(item.get("timestamp", ""))
                    row[2].text = item.get("source", "")
                    row[3].text = item.get("description", "")[:100]
                    row[4].text = item.get("risk_level", "")

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()

            logger.info(f"Generated DOCX: {len(docx_bytes)} bytes")
            return docx_bytes

        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            return None


class DigitalSigner:
    """Handle digital signatures for reports using cryptography library."""

    def __init__(self):
        self._crypto = None

    def _load_crypto(self):
        """Lazy load cryptography library."""
        if self._crypto is None:
            try:
                from cryptography.hazmat.primitives import hashes, serialization
                from cryptography.hazmat.primitives.asymmetric import rsa, padding
                from cryptography.hazmat.backends import default_backend
                from cryptography.x509 import load_pem_x509_certificate
                self._crypto = {
                    "hashes": hashes,
                    "serialization": serialization,
                    "rsa": rsa,
                    "padding": padding,
                    "default_backend": default_backend,
                    "load_pem_x509_certificate": load_pem_x509_certificate
                }
                logger.info("cryptography loaded successfully")
            except ImportError:
                logger.warning("cryptography not available, digital signatures disabled")
                self._crypto = {}
        return self._crypto

    def generate_key_pair(self, key_size: int = 2048) -> tuple:
        """Generate RSA key pair for signing."""
        crypto = self._load_crypto()
        if not crypto:
            return None, None

        try:
            private_key = crypto["rsa"].generate_private_key(
                public_exponent=65537,
                key_size=key_size,
                backend=crypto["default_backend"]()
            )

            private_pem = private_key.private_bytes(
                encoding=crypto["serialization"].Encoding.PEM,
                format=crypto["serialization"].PrivateFormat.PKCS8,
                encryption_algorithm=crypto["serialization"].NoEncryption()
            )

            public_pem = private_key.public_key().public_bytes(
                encoding=crypto["serialization"].Encoding.PEM,
                format=crypto["serialization"].PublicFormat.SubjectPublicKeyInfo
            )

            return private_pem, public_pem

        except Exception as e:
            logger.error(f"Key generation failed: {e}")
            return None, None

    def sign_document(self, document_bytes: bytes, private_key_pem: bytes) -> Optional[str]:
        """Sign document and return base64-encoded signature."""
        crypto = self._load_crypto()
        if not crypto:
            return None

        try:
            private_key = crypto["serialization"].load_pem_private_key(
                private_key_pem,
                password=None,
                backend=crypto["default_backend"]()
            )

            signature = private_key.sign(
                document_bytes,
                crypto["padding"].PSS(
                    mgf=crypto["padding"].MGF1(crypto["hashes"].SHA256()),
                    salt_length=crypto["padding"].PSS.MAX_LENGTH
                ),
                crypto["hashes"].SHA256()
            )

            return base64.b64encode(signature).decode("utf-8")

        except Exception as e:
            logger.error(f"Document signing failed: {e}")
            return None

    def verify_signature(self, document_bytes: bytes, signature_b64: str, public_key_pem: bytes) -> bool:
        """Verify document signature."""
        crypto = self._load_crypto()
        if not crypto:
            return False

        try:
            public_key = crypto["serialization"].load_pem_public_key(
                public_key_pem,
                backend=crypto["default_backend"]()
            )

            signature = base64.b64decode(signature_b64)

            public_key.verify(
                signature,
                document_bytes,
                crypto["padding"].PSS(
                    mgf=crypto["padding"].MGF1(crypto["hashes"].SHA256()),
                    salt_length=crypto["padding"].PSS.MAX_LENGTH
                ),
                crypto["hashes"].SHA256()
            )

            return True

        except Exception as e:
            logger.warning(f"Signature verification failed: {e}")
            return False

    def create_signature_info(
        self,
        signer_name: str,
        signer_title: str,
        signer_organization: str,
        document_bytes: bytes,
        private_key_pem: Optional[bytes] = None
    ) -> SignatureInfo:
        """Create signature info with optional cryptographic signature."""

        # Calculate document hash
        doc_hash = hashlib.sha256(document_bytes).hexdigest()

        # Sign if key provided
        signature_hash = None
        if private_key_pem:
            signature_hash = self.sign_document(document_bytes, private_key_pem)

        return SignatureInfo(
            signer_name=signer_name,
            signer_title=signer_title,
            signer_organization=signer_organization,
            signature_date=datetime.utcnow(),
            signature_hash=signature_hash,
            certificate_serial=doc_hash[:16].upper()
        )


class ReportEngine:
    """Main report generation engine."""

    def __init__(self):
        self.pdf_generator = PDFGenerator()
        self.docx_generator = DOCXGenerator()
        self.digital_signer = DigitalSigner()

    def generate_report(
        self,
        case_id: str,
        report_type: ReportType,
        format: ReportFormat,
        language: ReportLanguage,
        case_info: Dict[str, Any],
        evidence_items: List[EvidenceItem],
        chain_of_custody: Optional[ChainOfCustody] = None,
        risk_assessment: Optional[Dict[str, Any]] = None,
        sign_report: bool = False,
        signer_info: Optional[Dict[str, str]] = None,
        private_key_pem: Optional[bytes] = None
    ) -> GeneratedReport:
        """Generate a complete forensic report."""

        template = ReportTemplate(language)
        report_id = f"RPT-{case_id}-{datetime.utcnow().strftime('%Y%m%d%H%M%S')}"

        # Build report data
        report_data = self._build_report_data(
            case_id=case_id,
            report_type=report_type,
            template=template,
            case_info=case_info,
            evidence_items=evidence_items,
            chain_of_custody=chain_of_custody,
            risk_assessment=risk_assessment
        )

        # Generate HTML content
        html_content = self._generate_html(report_data, template, report_type)

        # Generate output based on format
        content = None
        page_count = 0

        if format == ReportFormat.PDF:
            css_styles = template.get_css_styles()
            content = self.pdf_generator.generate(html_content, css_styles)
            page_count = self.pdf_generator.get_page_count(html_content, css_styles)

        elif format == ReportFormat.DOCX:
            content = self.docx_generator.generate(report_data, template)
            page_count = len(report_data.get("sections", [])) + 2  # Estimate

        elif format == ReportFormat.HTML:
            full_html = f"""
            <!DOCTYPE html>
            <html lang="{language.value}">
            <head>
                <meta charset="UTF-8">
                <title>{report_data.get('title', 'Report')}</title>
                <style>{template.get_css_styles()}</style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            content = full_html.encode("utf-8")
            page_count = 1

        elif format == ReportFormat.JSON:
            # Serialize evidence items
            serialized_evidence = []
            for item in evidence_items:
                serialized_evidence.append({
                    "item_id": item.item_id,
                    "item_type": item.item_type,
                    "source": item.source,
                    "timestamp": item.timestamp.isoformat() if item.timestamp else None,
                    "description": item.description,
                    "hash_value": item.hash_value,
                    "metadata": item.metadata,
                    "risk_indicators": item.risk_indicators,
                    "relevance_score": item.relevance_score
                })

            json_data = {
                "report_id": report_id,
                "report_type": report_type.value,
                "case_id": case_id,
                "case_info": case_info,
                "evidence_items": serialized_evidence,
                "risk_assessment": risk_assessment,
                "generated_at": datetime.utcnow().isoformat()
            }
            content = json.dumps(json_data, indent=2, ensure_ascii=False).encode("utf-8")
            page_count = 1

        if content is None:
            # Fallback to HTML
            content = html_content.encode("utf-8")
            page_count = 1

        # Calculate content hash
        content_hash = hashlib.sha256(content).hexdigest()

        # Sign if requested
        signature = None
        if sign_report and signer_info:
            signature = self.digital_signer.create_signature_info(
                signer_name=signer_info.get("name", "Unknown"),
                signer_title=signer_info.get("title", "Forensic Analyst"),
                signer_organization=signer_info.get("organization", "SafeChild"),
                document_bytes=content,
                private_key_pem=private_key_pem
            )

        return GeneratedReport(
            report_id=report_id,
            report_type=report_type,
            format=format,
            language=language,
            title=report_data.get("title", template.t("report_title")),
            generated_at=datetime.utcnow(),
            content=content,
            content_hash=content_hash,
            page_count=page_count,
            signature=signature,
            metadata={
                "case_id": case_id,
                "evidence_count": len(evidence_items),
                "has_chain_of_custody": chain_of_custody is not None,
                "has_risk_assessment": risk_assessment is not None
            }
        )

    def _build_report_data(
        self,
        case_id: str,
        report_type: ReportType,
        template: ReportTemplate,
        case_info: Dict[str, Any],
        evidence_items: List[EvidenceItem],
        chain_of_custody: Optional[ChainOfCustody],
        risk_assessment: Optional[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Build structured report data."""

        # Title based on report type
        titles = {
            ReportType.FULL_ANALYSIS: template.t("report_title"),
            ReportType.EXECUTIVE_SUMMARY: template.t("executive_summary"),
            ReportType.EVIDENCE_INDEX: template.t("evidence_overview"),
            ReportType.CHAIN_OF_CUSTODY: template.t("chain_of_custody"),
            ReportType.RISK_ASSESSMENT: template.t("risk_assessment"),
            ReportType.EXPERT_WITNESS: "Expert Witness Statement",
            ReportType.ALIENATION_REPORT: "Parental Alienation Analysis",
            ReportType.TIMELINE: "Evidence Timeline"
        }

        data = {
            "title": titles.get(report_type, template.t("report_title")),
            "case_id": case_id,
            "case_info": case_info,
            "generated_at": datetime.utcnow(),
            "sections": [],
            "evidence_items": []
        }

        # Add evidence items as dicts
        for item in evidence_items:
            risk_level = "LOW"
            if item.relevance_score >= 0.7:
                risk_level = "HIGH"
            elif item.relevance_score >= 0.4:
                risk_level = "MEDIUM"

            data["evidence_items"].append({
                "item_id": item.item_id,
                "item_type": item.item_type,
                "source": item.source,
                "timestamp": item.timestamp.isoformat() if item.timestamp else "",
                "description": item.description,
                "hash_value": item.hash_value,
                "risk_level": template.t(risk_level.lower()),
                "risk_indicators": item.risk_indicators
            })

        # Add chain of custody if available
        if chain_of_custody:
            data["chain_of_custody"] = {
                "evidence_id": chain_of_custody.evidence_id,
                "acquisition_date": chain_of_custody.acquisition_date.isoformat(),
                "acquisition_method": chain_of_custody.acquisition_method,
                "acquired_by": chain_of_custody.acquired_by,
                "device_info": chain_of_custody.device_info,
                "hash_original": chain_of_custody.hash_original,
                "hash_algorithm": chain_of_custody.hash_algorithm,
                "custody_log": chain_of_custody.custody_log
            }

        # Add risk assessment if available
        if risk_assessment:
            data["risk_assessment"] = risk_assessment

        return data

    def _generate_html(
        self,
        report_data: Dict[str, Any],
        template: ReportTemplate,
        report_type: ReportType
    ) -> str:
        """Generate HTML content for the report."""

        html_parts = []

        # Header
        html_parts.append(f"""
        <div class="header">
            <h1>{report_data.get('title', template.t('report_title'))}</h1>
            <div class="confidential-banner">{template.t('confidential')}</div>
        </div>
        """)

        # Case info
        case_info = report_data.get("case_info", {})
        html_parts.append(f"""
        <div class="case-info">
            <table>
                <tr>
                    <td>{template.t('case_number')}:</td>
                    <td>{report_data.get('case_id', '')}</td>
                </tr>
                <tr>
                    <td>{template.t('date')}:</td>
                    <td>{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}</td>
                </tr>
        """)

        for key, value in case_info.items():
            html_parts.append(f"""
                <tr>
                    <td>{key}:</td>
                    <td>{value}</td>
                </tr>
            """)

        html_parts.append("</table></div>")

        # Executive summary for full reports
        if report_type == ReportType.FULL_ANALYSIS:
            evidence_count = len(report_data.get("evidence_items", []))
            high_risk = sum(1 for e in report_data.get("evidence_items", [])
                          if e.get("risk_level") == template.t("high"))

            html_parts.append(f"""
            <h2>{template.t('executive_summary')}</h2>
            <p>This forensic analysis report contains <strong>{evidence_count}</strong> evidence items,
            of which <strong class="risk-high">{high_risk}</strong> are classified as high risk.</p>
            """)

        # Evidence table
        if report_data.get("evidence_items"):
            html_parts.append(f"""
            <h2>{template.t('evidence_overview')}</h2>
            <table class="evidence-table">
                <thead>
                    <tr>
                        <th>ID</th>
                        <th>{template.t('timestamp')}</th>
                        <th>{template.t('source')}</th>
                        <th>{template.t('description')}</th>
                        <th>{template.t('risk_level')}</th>
                    </tr>
                </thead>
                <tbody>
            """)

            for item in report_data["evidence_items"]:
                risk_class = "risk-low"
                if item.get("risk_level") == template.t("high"):
                    risk_class = "risk-high"
                elif item.get("risk_level") == template.t("medium"):
                    risk_class = "risk-medium"

                html_parts.append(f"""
                    <tr>
                        <td>{item.get('item_id', '')}</td>
                        <td>{item.get('timestamp', '')}</td>
                        <td>{item.get('source', '')}</td>
                        <td>{item.get('description', '')[:100]}...</td>
                        <td class="{risk_class}">{item.get('risk_level', '')}</td>
                    </tr>
                """)

            html_parts.append("</tbody></table>")

        # Chain of custody
        if "chain_of_custody" in report_data:
            coc = report_data["chain_of_custody"]
            html_parts.append(f"""
            <div class="chain-of-custody">
                <h3>{template.t('chain_of_custody')}</h3>
                <p><strong>Evidence ID:</strong> {coc.get('evidence_id', '')}</p>
                <p><strong>Acquisition Date:</strong> {coc.get('acquisition_date', '')}</p>
                <p><strong>Method:</strong> {coc.get('acquisition_method', '')}</p>
                <p><strong>Acquired By:</strong> {coc.get('acquired_by', '')}</p>
                <div class="hash-verification">
                    <strong>{template.t('hash_verification')} ({coc.get('hash_algorithm', 'SHA-256')}):</strong><br>
                    {coc.get('hash_original', '')}
                </div>
            </div>
            """)

        # Risk assessment section
        if "risk_assessment" in report_data:
            ra = report_data["risk_assessment"]
            html_parts.append(f"""
            <h2>{template.t('risk_assessment')}</h2>
            <p><strong>Overall Risk Level:</strong>
                <span class="risk-{ra.get('overall_level', 'low').lower()}">{ra.get('overall_level', 'LOW')}</span>
            </p>
            <p><strong>Risk Score:</strong> {ra.get('risk_score', 0)}/10</p>
            """)

            if ra.get("indicators"):
                html_parts.append("<h3>Risk Indicators:</h3><ul>")
                for indicator in ra["indicators"]:
                    html_parts.append(f"<li>{indicator}</li>")
                html_parts.append("</ul>")

        return "\n".join(html_parts)


# Convenience function
def generate_forensic_report(
    case_id: str,
    report_type: str,
    format: str,
    language: str,
    case_info: Dict[str, Any],
    evidence_items: List[Dict[str, Any]],
    **kwargs
) -> GeneratedReport:
    """Convenience function to generate a forensic report."""

    engine = ReportEngine()

    # Convert evidence items to EvidenceItem objects
    evidence_objects = []
    for item in evidence_items:
        evidence_objects.append(EvidenceItem(
            item_id=item.get("item_id", ""),
            item_type=item.get("item_type", "unknown"),
            source=item.get("source", "unknown"),
            timestamp=item.get("timestamp") if isinstance(item.get("timestamp"), datetime)
                      else datetime.fromisoformat(item.get("timestamp", datetime.utcnow().isoformat())),
            description=item.get("description", ""),
            hash_value=item.get("hash_value", ""),
            metadata=item.get("metadata", {}),
            risk_indicators=item.get("risk_indicators", []),
            relevance_score=item.get("relevance_score", 0.0)
        ))

    return engine.generate_report(
        case_id=case_id,
        report_type=ReportType(report_type),
        format=ReportFormat(format),
        language=ReportLanguage(language),
        case_info=case_info,
        evidence_items=evidence_objects,
        **kwargs
    )
